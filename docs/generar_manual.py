from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x0d, 0x1b, 0x2e)
    hs.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    hs.paragraph_format.space_after = Pt(6)
    if level == 1: hs.font.size = Pt(22)
    elif level == 2: hs.font.size = Pt(16)
    else: hs.font.size = Pt(13)

def add_para(text, bold=False, italic=False, size=None, spacing=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    if size: r.font.size = Pt(size)
    if spacing: p.paragraph_format.space_after = Pt(spacing)
    return p

def bullet(text, bold_prefix=None, level=0):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix); r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    if level > 0: p.paragraph_format.left_indent = Inches(0.25*level)

def info_box(title, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(f"> {title}: "); r.bold = True; r.font.size = Pt(10)
    r2 = p.add_run(text); r2.italic = True; r2.font.size = Pt(10)

def add_table(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Light Grid Accent 1'
    for i,h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            for r in p.runs: r.bold = True; r.font.size = Pt(10)
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(10)

# ── PORTADA ──
for _ in range(6): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('MANUAL DE USUARIO'); r.bold = True; r.font.size = Pt(32); r.font.color.rgb = RGBColor(0x0d,0x1b,0x2e)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('SEE Ferreteria — Sistema de Gestion de Recambios'); r.font.size = Pt(18); r.font.color.rgb = RGBColor(0x2a,0x50,0x80)
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Version 1.0 — Julio 2026'); r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x66,0x66,0x66)
doc.add_page_break()

# ── INDICE ──
doc.add_heading('Indice', level=1)
for item in ['1. Introduccion','2. Inicio de sesion','3. El almacen — Vista general','4. Panel detallado','5. Recambios','6. Pedidos','7. Usuarios (solo administradores)','8. Solucion de problemas']:
    doc.add_paragraph(item).runs[0].font.size = Pt(12)
doc.add_page_break()

# ── 1. INTRODUCCION ──
doc.add_heading('1. Introduccion', level=1)
add_para('SEE Ferreteria es una aplicacion web disenada para la gestion integral de recambios en una ferreteria industrial. Su objetivo principal es organizar el almacen de forma visual —cada recambio tiene una ubicacion exacta dentro de paneles con cubetas— y facilitar la creacion de pedidos de reposicion, solicitudes personalizadas y urgencias.')
add_para('La aplicacion esta pensada para distintos perfiles dentro de la empresa: desde operarios que consultan ubicaciones y crean pedidos, hasta administradores que gestionan usuarios, recambios y paneles.', spacing=8)

doc.add_heading('Roles de usuario', level=2)
add_table(['Rol','Descripcion'],[['Admin','Acceso completo: gestionar recambios, paneles, pedidos y usuarios.'],['User','Puede ver recambios y paneles, crear pedidos.'],['Operario','Similar a user, con permisos ajustables por el administrador.'],['Viewer','Solo lectura: consulta recambios y paneles, no puede crear pedidos.']])
add_para('Cada rol tiene permisos configurables para las secciones de pedidos y recambios (crear, ver, editar, eliminar). El administrador puede ajustarlos desde la pagina de Usuarios.', spacing=8)

# ── 2. INICIO DE SESION ──
doc.add_heading('2. Inicio de sesion', level=1)
add_para('Al abrir la aplicacion por primera vez, te aparece la pantalla de inicio de sesion. Desde aqui puedes acceder de tres formas distintas.', spacing=8)

doc.add_heading('2.1 Acceder con usuario y contrasena', level=2)
add_para('Es el metodo principal de acceso. Sigue estos pasos:')
bullet('Introduce tu nombre de usuario en el campo correspondiente.', bold_prefix='Usuario: ')
bullet('Escribe tu contrasena.', bold_prefix='Contrasena: ')
bullet('Haz clic en el boton Acceder.', bold_prefix='')
add_para('Si las credenciales son correctas, entras directamente al almacen. Si algo falla, veras un mensaje de error en rojo.', spacing=8)

doc.add_heading('2.2 Iniciar sesion con Microsoft', level=2)
add_para('Si tu empresa usa Microsoft 365, puedes iniciar sesion con tu cuenta corporativa:')
bullet('En la pantalla de login, haz clic en Iniciar con Microsoft.')
bullet('Seras redirigido a la pagina de inicio de sesion de Microsoft.')
bullet('Introduce tu correo corporativo y contrasena.')
bullet('Tras la autenticacion, volveras automaticamente a la aplicacion.')
add_para('La primera vez que accedas con Microsoft se creara una cuenta automaticamente.', spacing=8)
info_box('Nota','Si el boton de Microsoft aparece desactivado con el mensaje "MSAL no esta configurado", significa que el administrador no ha habilitado esta opcion. Usa el formulario de usuario y contrasena.')

doc.add_heading('2.3 Registrarse', level=2)
bullet('Haz clic en "?No tienes cuenta? Registrate".')
bullet('Rellena los campos: Usuario, Nombre, Contrasena y Confirmar contrasena.')
bullet('Haz clic en Registrarse.')
add_para('Una vez completado el registro, iniciaras sesion automaticamente.', spacing=8)

doc.add_heading('2.4 Cerrar sesion', level=2)
add_para('Para salir, haz clic en el boton Salir en la esquina superior derecha del navbar.', spacing=8)

# ── 3. ALMACEN ──
doc.add_heading('3. El almacen — Vista general', level=1)
add_para('Al entrar en la aplicacion, lo primero que ves es la vista general del almacen. En ella aparecen todos los paneles (estanterias) como tarjetas horizontales que puedes desplazar con el raton o trackpad.')

doc.add_heading('3.1 ?Que es un panel?', level=2)
add_para('Un panel es una estanteria o zona del almacen identificada con un codigo (A1, A2, ..., A9, etc.). Cada panel tiene un numero fijo de cubetas (huecos):')
bullet('Paneles A1 a A5: 4 columnas x 8 filas = 32 cubetas.')
bullet('Paneles A6 a A9: 5 columnas x 10 filas = 50 cubetas.')
bullet('Otros paneles: 6 columnas x 15 filas = 90 cubetas.')

doc.add_heading('3.2 Leer una tarjeta de panel', level=2)
add_para('Cada tarjeta de panel muestra:')
bullet('Codigo del panel en azul (ej: A1).', bold_prefix='Titulo: ')
bullet('Nombre de la familia asignada o la mas comun entre sus recambios.', bold_prefix='Familia: ')
bullet('Cuadricula de cubetas:', bold_prefix='Cada cuadrito: ')
bullet('Color azul -> hay un recambio ocupando esa posicion.', level=1)
bullet('Gris oscuro -> la cubeta esta vacia.', level=1)
bullet('Icono [caja] -> hay recambio pero sin foto asignada.', level=1)

doc.add_heading('3.3 Asignar una familia a un panel', level=2)
bullet('Haz clic sobre el nombre de la familia en la tarjeta del panel.')
bullet('Selecciona la familia del desplegable y haz clic en Listo.')

# ── 4. PANEL DETALLADO ──
doc.add_heading('4. Panel detallado', level=1)
add_para('Haz clic en cualquier panel para ver su contenido en detalle. Cada celda muestra:')
bullet('Foto del recambio o icono [caja] si no tiene imagen.')
bullet('Referencia CMH en azul y referencia del cliente en naranja.')
bullet('Nombre del recambio y posicion (columna/fila).')

doc.add_heading('4.1 Acciones disponibles', level=2)
add_table(['Boton','Funcion'],[['Mostrar/Ocultar ocultos','Alterna la visualizacion de recambios ocultos (borde rojo).'],['Intercambiar/Mover','Activa el modo de intercambio o movimiento entre cubetas.']])

doc.add_heading('4.2 Intercambiar dos recambios', level=2)
bullet('Activa Intercambiar/Mover, haz clic en el primer recambio (se resalta en amarillo).')
bullet('Haz clic en el segundo recambio del mismo panel.')
bullet('Confirma el intercambio en la ventana emergente.')

doc.add_heading('4.3 Mover a otro panel', level=2)
bullet('Activa Intercambiar/Mover y selecciona un recambio.')
bullet('Haz clic en Mover a otro panel, elige el panel de destino y una cubeta vacia.')

# ── 5. RECAMBIOS ──
doc.add_heading('5. Recambios', level=1)
add_para('Cada recambio tiene una ubicacion fisica (panel, columna, fila) y una ficha con toda su informacion.', spacing=8)

doc.add_heading('5.1 Ficha tecnica', level=2)
add_para('Haz clic en cualquier recambio para abrir su ficha. Tiene tres pestanas:')
add_table(['Pestana','Descripcion'],[['Info','Datos generales, referencias, ubicacion, pedidos pendientes.'],['Historial','Pedidos creados para este recambio, del mas reciente al mas antiguo.'],['Nuevo Pedido','Crear pedido de reposicion, solicitud personalizada o urgente.']])
add_para('En Info, los administradores ven botones para Editar, Ocultar/Mostrar y Eliminar.', spacing=8)

doc.add_heading('5.2 Crear nuevo recambio', level=2)
add_para('Desde el Almacen, haz clic en + Nuevo Recambio. Campos obligatorios (*):')
add_table(['Campo','Descripcion','Ejemplo'],[['Ref. CMH','Codigo identificador unico','CMH-00123'],['Nombre','Nombre descriptivo','Tornillo M8x30'],['Panel','Estanteria','A1'],['Columna','Columna en el panel','3'],['Fila','Fila en el panel','5'],['Familia','Categoria','Tornilleria'],['Unidad de embalaje','Tipo de empaquetado','Caja de 10']])
add_para('Campos opcionales:', bold=True)
add_table(['Campo','Descripcion'],[['Ref. Cliente','Referencia del cliente'],['Codigo','Codigo interno'],['Marca','Fabricante'],['Descripcion','Notas detalladas'],['Metrica','Dimension tecnica (M8x30, 1/2")'],['Plazo de entrega','Plazo estimado (48h, 1 semana)'],['N Reposicion','Cantidad por defecto para pedidos automaticos']])
info_box('Importante','El boton de guardar se deshabilita hasta que todos los campos obligatorios esten correctos.')

doc.add_heading('5.3 Subir o quitar imagen', level=2)
bullet('Haz clic en Seleccionar imagen y elige un archivo (JPG, PNG, GIF, WebP; max 5 MB).', bold_prefix='Subir: ')
bullet('Quitar: haz clic en X Quitar si el recambio ya tiene imagen y quieres eliminarla.')

doc.add_heading('5.4 Buscar recambios', level=2)
add_para('Usa la barra de busqueda en el navbar. Escribe texto (referencia, nombre, marca) y aparecen sugerencias en tiempo real.')

doc.add_heading('5.5 Buscar por QR', level=2)
add_para('Haz clic en el boton QR del navbar, apunta la camara al codigo QR y se abrira la ficha del recambio.')

doc.add_heading('5.6 Ocultar o eliminar (solo admin)', level=2)
bullet('Ocultar: el recambio deja de verse salvo con "Mostrar ocultos" activado.', bold_prefix='Ocultar: ')
bullet('Mostrar: revierte la accion.', bold_prefix='Mostrar: ')
bullet('Eliminar: borrado permanente, requiere confirmacion.', bold_prefix='Eliminar: ')

# ── 6. PEDIDOS ──
doc.add_heading('6. Pedidos', level=1)
add_para('Los pedidos permiten solicitar recambios. Estados: Solicitado -> Pedido realizado -> Pedido recibido -> Finalizado.', spacing=8)

doc.add_heading('6.1 Crear pedido', level=2)
add_para('Desde la ficha tecnica, pestana Nuevo Pedido. Tres opciones:', spacing=4)

add_para('Automatico (Reposicion)', bold=True)
add_para('Usa la cantidad configurada en N Reposicion del recambio. No requiere campos adicionales.')
add_para('Si no tiene numero de reposicion configurado, el boton se deshabilita. Edita el recambio o usa otra modalidad.', spacing=6)

add_para('Personalizado (Solicitud)', bold=True)
bullet('N paquetes: cantidad (se multiplica por la unidad de embalaje).')
bullet('Fecha deseada de entrega.')
add_para('Ambos campos son obligatorios.', spacing=6)

add_para('Urgente (Solicitud Express)', bold=True)
add_para('Similar a Personalizado pero se marca como prioritario. La cantidad es opcional (si se omite, usa nReposicion).', spacing=6)

doc.add_heading('6.2 Pagina de pedidos', level=2)
add_para('En el navbar, haz clic en Pedidos. Filtros disponibles: busqueda por texto, tipo (Reposicion/Solicitud/Urgente), rango de fechas, orden, y opcion de ver finalizados.')
add_para('Los pedidos urgentes se muestran con borde rojo y etiqueta URGENTE.', spacing=8)

# ── 7. USUARIOS ──
doc.add_heading('7. Usuarios (solo administradores)', level=1)
add_para('Los administradores ven el enlace Usuarios en el navbar. Permite:')
bullet('Listar todos los usuarios con nombre, usuario, rol y estado.')
bullet('Crear usuario: nombre, usuario, rol. Opcion de generar contrasena aleatoria.')
bullet('Editar: cambiar nombre, rol, permisos detallados (crear/ver/editar/eliminar para pedidos y recambios) y estado activo/inactivo.')
bullet('Activar/desactivar: un usuario inactivo no puede iniciar sesion.')

# ── 8. SOLUCION DE PROBLEMAS ──
doc.add_heading('8. Solucion de problemas', level=1)

add_table(['Problema','Solucion'],[
    ['Error de conexion','Verifica tu conexion a internet. Si persiste, contacta al administrador.'],
    ['No se cargan las imagenes', 'Si ves [caja] el recambio no tiene foto. Si ves icono roto, la URL puede haber expirado.'],
    ['No puedo crear pedido automatico','El recambio debe tener un N Reposicion configurado. Editalo y establecelo.'],
    ['La sesion se cierra sola','Tiene duracion limitada (normalmente 8h). Inicia sesion de nuevo.'],
    ['MSAL no configurado','Usa usuario/contrasena. Solicita al admin que configure las variables de Azure AD.'],
    ['Error al subir imagen','Maximo 5 MB. Formatos: JPG, PNG, GIF, WebP.'],
])

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('--- Fin del manual ---'); r.italic = True; r.font.color.rgb = RGBColor(0x99,0x99,0x99)

doc.save('docs/MANUAL.docx')
print("Word document created: docs/MANUAL.docx")
